import os
import re
import csv
import json
import logging
from typing import List, Optional, Dict, Any
from pypdf import PdfReader
from docx import Document

from eightfold.models import RawRecord, Location, ExperienceEntry, EducationEntry
from eightfold.pipeline.normalize import SKILLS_MAP, SKILLS_ALIASES

logger = logging.getLogger("eightfold.extractor")

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file using pypdf."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        logger.warning(f"Error reading PDF {file_path}: {e}")
        raise e
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        logger.warning(f"Error reading DOCX {file_path}: {e}")
        raise e
    return text

def extract_skills_from_text(text: str) -> List[str]:
    """Helper to scan text for skills in the canonical skills map."""
    found_skills = set()
    text_lower = text.lower()
    for alias, canonical in SKILLS_MAP.items():
        # Use word boundaries for aliases to avoid partial matches (e.g. 'c' matching inside 'cat')
        # Special characters like c++ need careful regex formatting
        escaped_alias = re.escape(alias)
        # Check if the alias is surrounded by word boundaries or punctuation/space
        pattern = rf'(?:\b|_){escaped_alias}(?:\b|_)'
        if alias == 'c++':
            pattern = r'(?:\b)c\+\+(?:\b|\s|$)'
        elif alias == 'node.js':
            pattern = r'(?:\b)node\.js(?:\b|\s|$)'
        
        if re.search(pattern, text_lower):
            found_skills.add(canonical)
    return list(found_skills)

def parse_resume_text(text: str, source_id: str) -> RawRecord:
    """
    Parses resume text using regex and heuristics.
    Heuristics:
    - Name: First non-empty line (if short and not containing common keywords)
    - Email: regex
    - Phone: regex
    - Skills: scanned against skills map
    - Experience: check years of experience (regex like 'X years of experience')
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Heuristic for name: first line that doesn't look like email or phone or section header
    full_name = None
    for line in lines[:5]:
        if len(line) < 50 and '@' not in line and not re.search(r'\d{5,}', line):
            if not any(header in line.lower() for header in ['resume', 'cv', 'experience', 'education', 'skills', 'summary']):
                full_name = line
                break
                
    # Email regex
    emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    # Deduplicate preserving order
    emails = list(dict.fromkeys([e.strip() for e in emails]))
    
    # Phone regex
    # Match standard format: +1-234-567-8901, (123) 456-7890, etc.
    phone_pattern = r'(?:\+?\d{1,4}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    phones = list(dict.fromkeys([p.strip() for p in phones]))
    
    # Skills scan
    skills = extract_skills_from_text(text)
    
    # Years of experience heuristic
    years_exp = None
    m_exp = re.search(r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?(?:\s+of)?\s+experience', text, re.IGNORECASE)
    if m_exp:
        years_exp = float(m_exp.group(1))
        
    # Extract candidate's experience list (heuristics)
    experience_list = []
    # Search for blocks of experience: looking for common structures like "Company: X", "Title: Y"
    # or simple parsing lines. In a resume, we might just look for job titles or roles.
    # To keep it robust, let's extract simple experience entries if we find sections
    # or just parse them if the resume contains simple lines.
    # For a general resume, let's do a basic regex check.
    # Let's search for "Work History" or "Experience" section and extract lines
    # matching "Company: X", "Title: Y" or similar.
    # If not found, we can parse standard formats.
    # Let's also look for a common date range pattern: (2018 - 2022) or (2018-05 to 2021-08)
    exp_matches = re.finditer(r'(?:Company|Employer|At):\s*(.*?)\n(?:Title|Role|Position):\s*(.*?)\n(?:Dates|Period):\s*(.*?)(?:\n|$)', text, re.IGNORECASE)
    for match in exp_matches:
        company = match.group(1).strip()
        title = match.group(2).strip()
        dates_raw = match.group(3).strip()
        
        # Split dates
        start, end = None, None
        date_parts = re.split(r'[-–to]+', dates_raw)
        if len(date_parts) >= 1:
            start = date_parts[0].strip()
        if len(date_parts) >= 2:
            end = date_parts[1].strip()
            
        experience_list.append(ExperienceEntry(
            company=company,
            title=title,
            start=start,
            end=end,
            summary=None
        ))

    # If no structured experience blocks found, check for simple lines in sections
    if not experience_list:
        # Check if we can find a company name and title via simple line match in experience section
        # (Usually "Job Title at Company Name (Date - Date)")
        exp_section_match = re.search(r'(?:Experience|Work History|Professional Experience):?\s*\n(.*?)(?:\n\n|\n[A-Z][a-z]+:|\Z)', text, re.IGNORECASE | re.DOTALL)
        if exp_section_match:
            exp_text = exp_section_match.group(1)
            # Find lines like "Software Engineer at Google (2018 - 2020)"
            matches = re.finditer(r'([A-Za-z\s]+)\s+at\s+([A-Za-z0-9\s]+)\s*\((\d{4}[-/]\d{2}|\d{4})\s*[-–to\s]+\s*(\d{4}[-/]\d{2}|\d{4}|present|current)\)', exp_text, re.IGNORECASE)
            for m in matches:
                title = m.group(1).strip()
                company = m.group(2).strip()
                start = m.group(3).strip()
                end = m.group(4).strip()
                experience_list.append(ExperienceEntry(
                    company=company,
                    title=title,
                    start=start,
                    end=end
                ))

    return RawRecord(
        source_id=source_id,
        source_type='resume',
        method='regex_extract',
        full_name=full_name,
        emails=emails,
        phones=phones,
        skills=skills,
        years_experience=years_exp,
        experience=experience_list
    )

def parse_recruiter_notes(text: str, source_id: str) -> RawRecord:
    """
    Parses recruiter notes (.txt) line-by-line using regexes.
    Format is expected to be free text with fields like:
    Name: John Doe
    Email: john@example.com
    Phone: +1 234 567 8900
    Company: Acme Corp
    Title: Senior Engineer
    Skills: Python, Django, React
    Notes: ...
    """
    full_name = None
    emails = []
    phones = []
    company = None
    title = None
    skills = []
    years_exp = None

    lines = text.split('\n')
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Parse fields
        m_name = re.match(r'^(?:Candidate\s+)?Name\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_name:
            full_name = m_name.group(1).strip()
            continue
            
        m_email = re.match(r'^Email\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_email:
            emails = [e.strip() for e in m_email.group(1).split(',') if e.strip()]
            continue
            
        m_phone = re.match(r'^Phone\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_phone:
            phones = [p.strip() for p in m_phone.group(1).split(',') if p.strip()]
            continue
            
        m_comp = re.match(r'^(?:Current\s+)?Company\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_comp:
            company = m_comp.group(1).strip()
            continue
            
        m_title = re.match(r'^(?:Current\s+)?Title\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_title:
            title = m_title.group(1).strip()
            continue
            
        m_skills = re.match(r'^Skills\s*:\s*(.*)$', line_clean, re.IGNORECASE)
        if m_skills:
            # Recruiter notes skills might be comma separated
            skills_raw = m_skills.group(1).split(',')
            for sk in skills_raw:
                sk_clean = sk.strip()
                if sk_clean:
                    skills.append(sk_clean)
            continue
            
        m_exp = re.search(r'(?:Years\s+of\s+)?Experience\s*:\s*(\d+(?:\.\d+)?)\s*years?', line_clean, re.IGNORECASE)
        if m_exp:
            years_exp = float(m_exp.group(1))
            continue

    # Fallback for name in header
    if not full_name and lines:
        for line in lines[:3]:
            m_hdr = re.match(r'^(?:Recruiter\s+)?Notes\s+(?:on|for|on\s+candidate)\s+(.*)$', line.strip(), re.IGNORECASE)
            if m_hdr:
                full_name = m_hdr.group(1).strip()
                break

    # Map current company/title to an experience entry if available
    experience_list = []
    if company or title:
        experience_list.append(ExperienceEntry(
            company=company,
            title=title,
            start=None,
            end=None,
            summary="Extracted from recruiter notes"
        ))

    # Fallback to scanning the whole note for skills if no explicit skills field
    if not skills:
        skills = extract_skills_from_text(text)

    # Search for emails/phones if not explicitly specified in headers
    if not emails:
        emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if not phones:
        phone_pattern = r'(?:\+?\d{1,4}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)

    return RawRecord(
        source_id=source_id,
        source_type='recruiter_notes',
        method='heuristic_parse',
        full_name=full_name,
        emails=list(dict.fromkeys(emails)),
        phones=list(dict.fromkeys(phones)),
        skills=skills,
        years_experience=years_exp,
        experience=experience_list
    )

def extract_from_file(file_path: str, source_type: str) -> List[RawRecord]:
    """
    Loads candidate records from a single file path based on source_type.
    Catches file load and parse errors, logs them, and returns an empty list, ensuring robustness.
    """
    basename = os.path.basename(file_path)
    records: List[RawRecord] = []

    try:
        if source_type == 'recruiter_csv':
            with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                reader = csv.DictReader(f)
                for idx, row in enumerate(reader):
                    # Check field aliases
                    name = row.get('name') or row.get('full_name') or row.get('candidate_name')
                    email = row.get('email') or row.get('email_address')
                    phone = row.get('phone') or row.get('telephone') or row.get('phone_number')
                    company = row.get('current_company') or row.get('company')
                    title = row.get('title') or row.get('current_title') or row.get('job_title')
                    years_exp = row.get('years_experience') or row.get('experience_years')
                    
                    emails = [email.strip()] if email and email.strip() else []
                    phones = [phone.strip()] if phone and phone.strip() else []
                    
                    exp = []
                    if company or title:
                        exp.append(ExperienceEntry(
                            company=company.strip() if company else None,
                            title=title.strip() if title else None,
                            start=None,
                            end=None,
                            summary="Current role from Recruiter CSV"
                        ))

                    y_exp = None
                    if years_exp:
                        try:
                            y_exp = float(years_exp.strip())
                        except ValueError:
                            pass
                            
                    records.append(RawRecord(
                        source_id=f"{basename}_row_{idx}",
                        source_type='recruiter_csv',
                        method='direct_field',
                        full_name=name.strip() if name else None,
                        emails=emails,
                        phones=phones,
                        experience=exp,
                        years_experience=y_exp
                    ))

        elif source_type == 'ats_json':
            with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                
            items = data if isinstance(data, list) else [data]
            for idx, item in enumerate(items):
                ats_id = item.get('ats_id') or f"index_{idx}"
                name = item.get('candidate_name') or item.get('full_name') or item.get('name')
                email = item.get('email_address') or item.get('email')
                phone = item.get('telephone') or item.get('phone')
                company = item.get('company_name') or item.get('company')
                title = item.get('job_title') or item.get('title')
                years_exp = item.get('years_experience') or item.get('experience')
                
                # Check for nested structures in ATS JSON
                location_obj = None
                loc_data = item.get('location')
                if isinstance(loc_data, dict):
                    location_obj = Location(
                        city=loc_data.get('city'),
                        region=loc_data.get('region'),
                        country=loc_data.get('country')
                    )
                elif isinstance(loc_data, str):
                    # Try to parse string location: e.g. "San Francisco, CA, US"
                    parts = [p.strip() for p in loc_data.split(',')]
                    if len(parts) == 3:
                        location_obj = Location(city=parts[0], region=parts[1], country=parts[2])
                    elif len(parts) == 2:
                        location_obj = Location(city=parts[0], country=parts[1])
                    else:
                        location_obj = Location(city=parts[0])

                skills = item.get('skills', [])
                if isinstance(skills, str):
                    skills = [s.strip() for s in skills.split(',') if s.strip()]

                # Experience list
                exp_list = []
                # Check if experience is in a list format
                raw_exp_list = item.get('experience')
                if isinstance(raw_exp_list, list):
                    for e in raw_exp_list:
                        if isinstance(e, dict):
                            exp_list.append(ExperienceEntry(
                                company=e.get('company'),
                                title=e.get('title'),
                                start=e.get('start'),
                                end=e.get('end'),
                                summary=e.get('summary')
                            ))
                elif company or title:
                    exp_list.append(ExperienceEntry(
                        company=company,
                        title=title,
                        start=None,
                        end=None
                    ))

                # Education list
                edu_list = []
                raw_edu_list = item.get('education')
                if isinstance(raw_edu_list, list):
                    for ed in raw_edu_list:
                        if isinstance(ed, dict):
                            try:
                                end_yr = int(ed.get('end_year')) if ed.get('end_year') else None
                            except ValueError:
                                end_yr = None
                            edu_list.append(EducationEntry(
                                institution=ed.get('institution'),
                                degree=ed.get('degree'),
                                field=ed.get('field'),
                                end_year=end_yr
                            ))

                emails = [email.strip()] if email and email.strip() else []
                phones = [phone.strip()] if phone and phone.strip() else []
                
                y_exp = None
                if years_exp is not None:
                    try:
                        y_exp = float(years_exp)
                    except ValueError:
                        pass

                records.append(RawRecord(
                    source_id=f"{basename}_ats_{ats_id}",
                    source_type='ats_json',
                    method='direct_field',
                    full_name=name.strip() if name else None,
                    emails=emails,
                    phones=phones,
                    location=location_obj,
                    skills=skills,
                    experience=exp_list,
                    education=edu_list,
                    years_experience=y_exp
                ))

        elif source_type == 'resume':
            text = ""
            ext = os.path.splitext(basename)[1].lower()
            if ext == '.pdf':
                text = extract_text_from_pdf(file_path)
            elif ext == '.docx':
                text = extract_text_from_docx(file_path)
            else:
                # Text resume fallback
                with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                    
            if text.strip():
                records.append(parse_resume_text(text, basename))
            else:
                logger.warning(f"Extracted resume text is empty for {file_path}")

        elif source_type == 'recruiter_notes':
            with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if text.strip():
                records.append(parse_recruiter_notes(text, basename))

    except Exception as e:
        logger.warning(f"Failed to extract candidate records from {file_path}: {e}")
        # Gracefully degrade by returning empty list of records

    return records
